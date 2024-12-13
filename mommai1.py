import os
import openai 
from openai import Client
from flask import Flask, request, jsonify, render_template
from docx import Document
from dotenv import load_dotenv
import math

app = Flask(__name__, template_folder='templates')

@app.route('/test')
def test():
    return "Flask is working!"

# Load environment variables from the .env file
load_dotenv()

# Set up the OpenAI API key (ensure you have it in the .env file or set it directly here)
api_key = os.getenv("OPENAI_API_KEY")  # Retrieve the API key from environment variables

# Check if the API key is not found
if not api_key:
    raise ValueError("API key is missing from environment variables.")

# Initialize the OpenAI client
client = Client(api_key=api_key)  # Pass the API key to the client

# Define the path to the folder containing doc files
folder_path = "./static/files"

# Global variable to store document content
documents_content = {}

def load_documents(folder_path):
    """
    Load and parse all .docx files in the specified folder.
    """
    docs = {}
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(folder_path, filename)
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            docs[filename] = content
    return docs

# Load documents at startup
documents_content = load_documents(folder_path)
if not documents_content:
    print("No documents found in the folder.")
else:
    print(f"Loaded documents: {', '.join(documents_content.keys())}")

# Function to list all .docx files in the folder
def list_documents_in_folder(folder):
    documents = []
    for file in os.listdir(folder):
        if file.endswith(".docx"):
            documents.append(file)
    return documents

# Function to read the content of a specific .docx file
def read_doc_file(file_path):
    doc = Document(file_path)
    content = "\n".join([paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()])
    return content

# List the documents
documents = list_documents_in_folder(folder_path)
print("Documents in the static folder:")
for i, doc in enumerate(documents, 1):
    print(f"{i}. {doc}")

# Ask the user to select a document number to view
file_index = None
while file_index is None:
    try:
        file_index = int(input("\nEnter the number of the document you want to view: ")) - 1
        if not (0 <= file_index < len(documents)):
            print("Invalid selection, please choose a valid number.")
            file_index = None  # Reset the variable if invalid input
    except ValueError:
        print("Invalid input, please enter a number.")

# If a valid file index is selected, read and display the content
file_name = documents[file_index]
file_path = os.path.join(folder_path, file_name)
document_content = read_doc_file(file_path)

print(f"\nContent of {file_name}:")
print(document_content)

# Define cosine similarity (pure Python implementation)
def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors without NumPy."""
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = math.sqrt(sum(a * a for a in vec1))
    norm2 = math.sqrt(sum(b * b for b in vec2))
    return dot_product / (norm1 * norm2)

# Function to chat with Theresa
def chat_with_theresa(user_input, documents):
    """
    Generate a response from Theresa using the loaded documents.
    """

    if user_input.lower() == "show docs":
        # List all available documents
        return f"I have access to: {', '.join(documents.keys())}"

    # Combine all document content for context
    document_references = "\n\n".join([f"### {name}:\n{content}" for name, content in documents_content.items()])

    # Construct the prompt
    prompt = f"""
    You are Theresa, the user's mother, a wise and nurturing figure. You have have access to and have written the following writings as a guide to your sons, the user.:
    {document_references}

    Respond to the user's queries based on this content. Provide advice and wisdom as Theresa would, referencing relevant content from the writings where appropriate. Ensure to let the user know where the advice is referenced in the writings.:
    User says: {user_input}
    """

    # Use the OpenAI client to create a chat completion based on the prompt
    try: 
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Ensure the model is available and supports this functionality
            messages=[
                {"role": "system" , "content": "You are Theresa, a concerned mother of the user, wise and nurturing; constantly referring to her writings."},
                {"role": "user", "content": user_input}],
            max_tokens=800,
            temperature=0.7,
        )
        
        # Corrected access: use dot notation to get the content of the response
        content = response.choices[0].message.content  # Corrected attribute access
        return content  # Only return the response content

    except Exception as e:
        return f"An error occurred: {e}"

# Chat loop
print(f"\nYou are now chatting with Theresa based on the content of {file_name}.")
while True:
    user_input = input("Esrom: ")
    if user_input.lower() == "1251":
        print("Goodbye!")
        break
    
    # Get the AI's response based on the document content
    response = chat_with_theresa(user_input, document_content)  # Only get the response, no need for unpacking
    print(f"Theresa: {response}")

# Flask route for serving the index.html page
@app.route('/')
def home():
    return render_template('index.html')

# Flask route for user chat
@app.route('/chat', methods=['POST'])
def chat():
    """
    Endpoint to interact with Theresa via API.
    Expects JSON input with a "user_input" key.
    """
    data = request.json
    user_input = data.get("user_input", "")
    if not user_input:
        return jsonify({"error": "No input provided."}), 400

    # Respond to user input using all document content
    response = chat_with_theresa(user_input, document_content)
    return jsonify({"response": response})

if __name__ == '__main__':
    mode = input("Enter mode (flask/terminal): ").strip().lower()
    if mode == 'flask':
        app.run(host='0.0.0.0', port=5001)
    elif mode == 'terminal':
        print("You are now chatting with Theresa. Type 'exit' to quit.")
        while True:
            user_input = input("You: ")
            if user_input.lower() == "exit":
                print("Goodbye!")
                break
            # Simulate Theresa's response
            print(f"Theresa: You said '{user_input}'")
    else:
        print("Invalid mode. Exiting.")

